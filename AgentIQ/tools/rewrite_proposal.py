from __future__ import annotations

import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(r"M:\semester4\AIL\project")
OUTPUT = PROJECT_ROOT / "proposal.docx"
BACKUP = PROJECT_ROOT / "proposal_before_agentiq_update.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
LIGHT_FILL = "F4F6F9"
HEADER_FILL = "E8EEF5"
BORDER = "D9E2EC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = width
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc: Document, text: str, size=22, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def add_meta_line(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r2 = p.add_run(value)
    r2.font.color.rgb = MUTED


def add_bullets(doc: Document, items) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc: Document, items) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_table(doc: Document, headers, rows, widths) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        set_cell_shading(hdr[i], HEADER_FILL)
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = DARK_BLUE
    for row_data in rows:
        row = table.add_row().cells
        for i, value in enumerate(row_data):
            row[i].text = value
    set_table_widths(table, widths)
    doc.add_paragraph()


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_FILL)
    set_cell_margins(cell, top=140, bottom=140, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def build_doc() -> Document:
    doc = Document()
    style_document(doc)

    add_title(doc, "AI Project Proposal", size=24)
    add_title(doc, "AgentIQ: AI Search Algorithm Simulator with Context-Aware Assistant", size=15, color=DARK_BLUE)
    add_meta_line(doc, "Session", "2024 - 2028")
    add_meta_line(doc, "Course", "CSC-202L Artificial Intelligence")
    add_meta_line(doc, "Submitted to", "Ms. Fiza Khalid")
    add_meta_line(doc, "Department", "Department of Computer Science, University of Engineering and Technology, Lahore, Pakistan")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Submitted by")
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    add_table(
        doc,
        ["Team Member", "Roll Number", "Project Role"],
        [
            ["Muntaha Fatima", "2024-CS-196", "Simulation core, dashboard progress, integration"],
            ["Wareesha Ameer Khan", "2024-CS-202", "Assistant logic, knowledge base, Gemini fallback"],
            ["Hania Bukhari", "2024-CS-220", "UI design, multi-page interface, layout polish"],
            ["Haram Naseeb", "2024-CS-230", "Problem environments, PEAS, domain modeling"],
        ],
        [Inches(1.75), Inches(1.35), Inches(3.15)],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)

    doc.add_heading("1. Project Overview", level=1)
    doc.add_paragraph(
        "AgentIQ is a client-side educational web application that demonstrates how classical "
        "Artificial Intelligence search algorithms solve real state-space problems. Instead of only "
        "showing final answers, the system exposes the complete reasoning process: the frontier, "
        "explored set, current node, path cost, heuristic values, pseudocode phase, and final solution path."
    )
    doc.add_paragraph(
        "The project is built with vanilla HTML5, CSS3, and JavaScript ES6 modules. It includes a "
        "visual simulation workspace, algorithm comparison tools, a curated AI knowledge base, optional "
        "Google Gemini fallback for unanswered questions, browser-based authentication, user-specific chat "
        "history, and user-specific learning progress tracking."
    )
    add_callout(
        doc,
        "Educational purpose",
        "The main goal is to convert abstract AI topics such as frontier management, admissible heuristics, PEAS, "
        "state transitions, and search optimality into visual, step-by-step learning experiences."
    )

    doc.add_heading("2. Project Objectives", level=1)
    add_numbered(
        doc,
        [
            "Provide an interactive platform for learning uninformed, informed, and local search algorithms.",
            "Visualize how nodes or states are selected, expanded, evaluated, and added to the frontier.",
            "Explain each supported problem using AI state-space components and PEAS descriptions.",
            "Support side-by-side algorithm comparison for cost, steps, and behavior.",
            "Assist learners through a context-aware tutor that uses an 80 percent knowledge-base confidence gate and falls back to Gemini when needed.",
            "Track each user's chat history and learning progress separately in the browser.",
        ],
    )

    doc.add_heading("3. Search Algorithms Covered", level=1)
    doc.add_paragraph(
        "The simulator covers six core AI search strategies. Each algorithm follows a common engine interface "
        "so that the same problem can be executed, paused, stepped, reset, and compared consistently."
    )
    add_table(
        doc,
        ["Category", "Algorithm", "How AgentIQ Demonstrates It"],
        [
            ["Uninformed", "Breadth-First Search (BFS)", "Uses FIFO expansion and demonstrates level-by-level search with shortest path behavior for equal costs."],
            ["Uninformed", "Depth-First Search (DFS)", "Uses LIFO expansion to show deep exploration, backtracking behavior, and memory-efficient traversal."],
            ["Uninformed", "Uniform Cost Search (UCS)", "Expands the lowest g(n) path cost first and demonstrates optimal weighted search."],
            ["Informed", "Greedy Best-First Search", "Uses h(n) only, showing fast heuristic guidance and why it may be suboptimal."],
            ["Informed", "A* Search", "Combines g(n) + h(n), showing f(n) calculations and optimality with suitable heuristics."],
            ["Local Search", "Hill Climbing", "Moves to the best neighbor and demonstrates local optima, plateaus, and optimization behavior."],
        ],
        [Inches(1.25), Inches(1.75), Inches(3.25)],
    )

    doc.add_heading("4. Problem Domains Covered", level=1)
    doc.add_paragraph(
        "AgentIQ includes twelve AI problem domains. Each domain defines an initial state, successor function, "
        "goal test, path cost or evaluation function, and, where appropriate, a heuristic."
    )
    add_table(
        doc,
        ["Problem", "Type", "What It Demonstrates"],
        [
            ["Romania Map", "Weighted graph", "City-to-city pathfinding, g(n), h(n), f(n), and shortest route planning."],
            ["8-Puzzle", "Sliding puzzle", "State transitions, tile movement, misplaced tiles, and Manhattan-style reasoning."],
            ["N-Queens", "Constraint problem", "Conflict minimization and local search behavior."],
            ["Vacuum World", "Agent environment", "Perception-action loop, cleaning goals, and PEAS modeling."],
            ["Wumpus World", "Logical/grid world", "Partially observable environment, percepts, and safe navigation."],
            ["Travelling Salesman Problem", "Optimization graph", "Route cost, city tour improvement, Greedy behavior, and Hill Climbing."],
            ["Sudoku", "Constraint satisfaction", "Grid constraints and conflict-based evaluation."],
            ["Maze Navigation", "Grid search", "Obstacle avoidance and pathfinding through cells."],
            ["Robot Path Planning", "Grid planning", "Autonomous movement, walls, start-goal navigation, and shortest path search."],
            ["Tower of Hanoi", "State-transition problem", "Legal moves, recursive structure, and state-space traversal."],
            ["Water Jug", "Classical planning", "Fill, empty, and pour operators for goal-state measurement."],
            ["Missionaries and Cannibals", "Constraint planning", "Safe transport rules and valid state generation."],
        ],
        [Inches(1.7), Inches(1.45), Inches(3.1)],
    )

    doc.add_heading("5. PEAS and Agent Modeling", level=1)
    doc.add_paragraph(
        "Every problem is presented with a PEAS description: Performance Measure, Environment, Actuators, "
        "and Sensors. This connects the visual simulations to the intelligent-agent framework taught in AI."
    )
    add_bullets(
        doc,
        [
            "Performance explains how success is measured, such as reaching a goal with low cost or cleaning all dirty cells.",
            "Environment defines the world the agent operates in, such as a grid, graph, board, cave, or puzzle state.",
            "Actuators describe the actions available to the agent, such as moving, sliding, pouring, or placing pieces.",
            "Sensors describe what information the agent can observe, such as current position, tile layout, percepts, or board conflicts.",
        ],
    )

    doc.add_heading("6. Interactive Simulation Workspace", level=1)
    doc.add_paragraph(
        "The simulation page is the main learning workspace. It uses a split layout with a large visual area on "
        "the left and controls, PEAS, metrics, path details, pseudocode, and the assistant panel on the right."
    )
    add_bullets(
        doc,
        [
            "Playback controls include Start, Pause, Resume, Step Forward, Step Backward, and Reset.",
            "A speed slider lets the learner slow down or accelerate the visualization.",
            "Live metrics show current step, expanded nodes, current node, frontier nodes, explored nodes, path cost, and solution path.",
            "The heuristic panel explains h(n) and f(n) values for informed algorithms such as Greedy and A*.",
            "The pseudocode panel highlights the active line, connecting code-level algorithm logic with visual state changes.",
            "Canvas renderers are used for graph/grid-style problems, while DOM renderers are used for tile or board-style problems.",
        ],
    )

    doc.add_heading("7. AgentIQ Assistant", level=1)
    doc.add_paragraph(
        "The AgentIQ Assistant is a context-aware tutor embedded into the simulator. It combines a curated knowledge "
        "base with optional Gemini AI. The assistant first checks whether the question has a high-confidence knowledge-base "
        "answer. If the KB confidence is at least 80 percent, it answers from the trusted local content. If no strong answer "
        "is found, it sends the current simulation context to Gemini."
    )
    add_bullets(
        doc,
        [
            "Knowledge-base answers cover AI definitions, PEAS, BFS, DFS, UCS, Greedy, A*, Hill Climbing, heuristics, complexity, and problem descriptions.",
            "Dynamic answers use the live simulation state to explain current nodes, g(n), h(n), f(n), path cost, frontier contents, and selected algorithm behavior.",
            "Gemini fallback is used for novel or detailed questions that the local knowledge base cannot confidently answer.",
            "The model selection logic avoids deprecated Gemini model names and can list available generateContent models for the configured API key.",
            "Suggested questions change according to the active problem and algorithm.",
            "Each logged-in user has a separate chat history and can clear their own chat from the assistant panel.",
        ],
    )

    doc.add_heading("8. User Accounts, Chat History, and Progress", level=1)
    doc.add_paragraph(
        "AgentIQ includes lightweight browser-based authentication for project demonstration. User accounts, sessions, "
        "chat history, preferences, and progress are stored locally in the browser using localStorage."
    )
    add_table(
        doc,
        ["Feature", "Implementation", "Purpose"],
        [
            ["User login/signup", "Mock authentication stored in localStorage", "Allows separate learning sessions for different students on the same browser."],
            ["User-scoped chat history", "Stored with a key based on the logged-in email", "Prevents one student's assistant conversation from mixing with another's."],
            ["Clear Chat", "Assistant panel control", "Lets the user reset their own chat history without affecting other users."],
            ["User-scoped progress", "Progress stored with a key based on the logged-in email", "Tracks completed problems and algorithms for each user separately."],
            ["Dashboard overview", "Progress cards and completed badges", "Shows completed problems, algorithms tried, completion percentage, and latest completion date."],
        ],
        [Inches(1.55), Inches(2.35), Inches(2.35)],
    )

    doc.add_heading("9. Multi-Page Application Structure", level=1)
    add_table(
        doc,
        ["Page / Module", "Responsibility"],
        [
            ["index.html", "Landing page introducing AgentIQ, features, workflow, and problem catalog."],
            ["login.html", "User signup and login flow for browser-based sessions."],
            ["dashboard.html", "Problem cards, filters, user progress overview, and completed badges."],
            ["algo-select.html", "Algorithm selection page that shows compatible algorithms for the chosen problem."],
            ["simulation.html", "Main visual simulator with controls, PEAS, metrics, pseudocode, and assistant."],
            ["comparison.html", "Side-by-side comparison of two algorithms on the same problem."],
            ["assets/js/algorithms", "Algorithm engines for BFS, DFS, UCS, Greedy, A*, and Hill Climbing."],
            ["assets/js/problems", "State-space problem models and render logic."],
            ["assets/js/assistant", "Knowledge base matcher, Gemini API integration, context builder, suggestions, and chat panel."],
        ],
        [Inches(2.0), Inches(4.25)],
    )

    doc.add_heading("10. Algorithm Comparison Feature", level=1)
    doc.add_paragraph(
        "The comparison page allows learners to run two algorithms on the same selected problem. This is useful for "
        "observing practical differences between search strategies, such as BFS versus DFS, UCS versus A*, or Greedy "
        "versus A*. The comparison uses isolated state objects so that each algorithm can run independently without "
        "overwriting the other algorithm's frontier, explored set, path, or metrics."
    )

    doc.add_heading("11. Technical Implementation", level=1)
    add_bullets(
        doc,
        [
            "Frontend stack: HTML5, CSS3, JavaScript ES6 modules.",
            "State management: shared simulation state object with snapshots for step-back history.",
            "Rendering: canvas renderer for spatial/graph problems and DOM renderer for board/tile problems.",
            "Storage: localStorage for mock accounts, sessions, theme, Gemini key configuration, chat history, and progress.",
            "AI assistant: local KB matcher, 80 percent confidence gate, context prompt builder, and Gemini generateContent fallback.",
            "Responsive UI: dashboard uses 3 problem cards per row on desktop, 2 on tablet, and 1 on mobile.",
        ],
    )

    doc.add_heading("12. Key Design Principles", level=1)
    add_bullets(
        doc,
        [
            "Educational clarity: every feature explains an AI concept instead of only showing an output.",
            "Modularity: algorithms, problems, renderers, assistant logic, and state handling are separated into modules.",
            "Traceability: step history, pseudocode highlighting, frontier display, and metrics make each algorithm decision observable.",
            "User personalization: each user has separate chat and progress records.",
            "Graceful fallback: the assistant answers from the curated KB when confident and switches to Gemini when a question is outside the KB.",
            "Visual readability: the dashboard, simulation panels, cards, and progress displays are designed for quick scanning.",
        ],
    )

    doc.add_heading("13. Expected Learning Outcomes", level=1)
    add_bullets(
        doc,
        [
            "Students can distinguish uninformed, informed, and local search strategies.",
            "Students can explain how frontier ordering changes the behavior of BFS, DFS, UCS, Greedy, and A*.",
            "Students can understand g(n), h(n), and f(n) through live examples.",
            "Students can connect AI agents to PEAS descriptions for multiple environments.",
            "Students can compare algorithms empirically using visual progress, cost, and path behavior.",
            "Students can ask context-aware questions and receive explanations connected to the current simulation state.",
        ],
    )

    doc.add_heading("14. Conclusion", level=1)
    doc.add_paragraph(
        "AgentIQ successfully combines AI theory, visualization, and an intelligent tutoring layer into one educational "
        "platform. The project covers classical search algorithms, twelve AI problem environments, PEAS analysis, "
        "live pseudocode tracing, comparison tools, user-specific learning records, and a hybrid KB/Gemini assistant. "
        "Together, these features make the project a complete and practical AI learning simulator for students."
    )

    return doc


def main() -> None:
    output = Path(sys.argv[1]) if len(sys.argv) > 1 else OUTPUT
    backup = output.with_name(f"{output.stem}_before_agentiq_update{output.suffix}")

    if output == OUTPUT and OUTPUT.exists() and not BACKUP.exists():
        shutil.copy2(OUTPUT, BACKUP)
        backup = BACKUP
    elif output.exists() and not backup.exists():
        shutil.copy2(output, backup)

    doc = build_doc()
    doc.save(output)
    print(f"Updated {output}")
    print(f"Backup at {backup}")
    print(f"Generated at {datetime.now().isoformat(timespec='seconds')}")


if __name__ == "__main__":
    main()
